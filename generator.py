from docx import Document
from docx.shared import Pt
from typing import List
from customtkinter import CTkLabel
from recording import Recording
from utils import remove_extra_whitespaces


class Generator:
    def __init__(self, info_label: CTkLabel) -> None:
        self.group: bool = False
        self.info_label = info_label
        self.time_interval_regex = r"(\b[0-2]?[0-9].[0-5][0-9]\b-\b[0-2]?[0-9].[0-5][0-9]\b)"
        self.time_regex = r"(\b[0-2]?[0-9].[0-5][0-9]\b)"
        self.uniq_groups: List[str] = []

    @staticmethod
    def __create_doc() -> Document:
        doc = Document()

        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(14)

        return doc

    def sort_key(self, record: Recording) -> tuple:
        record_list = record.get_record()
        split_data = record_list[0][:5].split('.')

        if self.group:
            return record_list[3], split_data[1], split_data[0]
        else:
            return split_data[1], split_data[0]

    def __output_results(self, items: List[Recording], doc: Document) -> None:
        items.sort(key=self.sort_key)

        for result_item in items:
            result_string = result_item.get_record()
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.line_spacing = 1
            paragraph.add_run(remove_extra_whitespaces(' '.join(result_string[:2]))).bold = True
            paragraph.add_run(' ' + remove_extra_whitespaces(' '.join(result_string[2:])))

    def generate(self, results: List[List[Recording]], search_filter: List[str], general_file: bool = False,
                 group: bool = False) -> None:
        self.group = group
        self.info_label.configure(text="Генерация...")

        if general_file:
            doc = self.__create_doc()

            for item in results:
                self.__output_results(item, doc)
                doc.add_paragraph().paragraph_format.line_spacing = 1

            doc.save(f"Расписание {' '.join(search_filter)}.docx")
        else:
            for item in results:
                doc = self.__create_doc()
                self.__output_results(item, doc)
                doc.save(f"Расписание {search_filter[results.index(item)]}.docx")

        self.info_label.configure(text="Генерация завершена")
